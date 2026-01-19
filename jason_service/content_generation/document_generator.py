import json
import logging
import os
import re
import time
from typing import Dict, List, Optional, Tuple, Any, Union
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import random

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - Word document generation disabled")

try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logging.warning("openpyxl not available - Excel spreadsheet generation disabled")

@dataclass
class DocumentSection:
    """Represents a section in a document."""
    title: str
    content: List[str]
    level: int = 1
    subsections: List['DocumentSection'] = None
    
    def __post_init__(self):
        if self.subsections is None:
            self.subsections = []

@dataclass
class DocumentStructure:
    """Complete document structure with metadata."""
    title: str
    subtitle: str
    author: str
    date: str
    sections: List[DocumentSection]
    document_type: str = "report"  # report, essay, proposal, manual
    style: str = "professional"
    total_words: int = 0

@dataclass
class ExcelSheet:
    """Represents a worksheet in an Excel workbook."""
    name: str
    data: List[List[Union[str, int, float]]]
    headers: List[str]
    chart_type: Optional[str] = None  # bar, line, pie
    calculations: List[Dict] = None
    
    def __post_init__(self):
        if self.calculations is None:
            self.calculations = []

@dataclass
class ExcelWorkbook:
    """Complete Excel workbook structure."""
    title: str
    sheets: List[ExcelSheet]
    theme: str = "professional"
    created_by: str = "JASON - AI Architect"
    created_date: str = datetime.now().strftime("%Y-%m-%d")

class DocumentContentExpander:
    """Advanced content expansion system for document generation."""
    
    def __init__(self):
        self.content_templates = self._load_content_templates()
        self.document_structures = self._load_document_structures()
        self.writing_styles = self._load_writing_styles()
    
    def _load_content_templates(self) -> Dict[str, Dict]:
        """Load templates for different types of documents."""
        return {
            'report': {
                'sections': [
                    'Executive Summary',
                    'Introduction',
                    'Methodology',
                    'Findings',
                    'Analysis',
                    'Recommendations',
                    'Conclusion',
                    'Appendices'
                ],
                'word_count_per_section': (150, 300),
                'tone': 'formal',
                'structure': 'hierarchical'
            },
            'essay': {
                'sections': [
                    'Introduction',
                    'Background',
                    'Main Arguments',
                    'Counterarguments',
                    'Conclusion'
                ],
                'word_count_per_section': (200, 400),
                'tone': 'academic',
                'structure': 'linear'
            },
            'proposal': {
                'sections': [
                    'Problem Statement',
                    'Proposed Solution',
                    'Implementation Plan',
                    'Budget',
                    'Timeline',
                    'Expected Outcomes',
                    'Risk Assessment'
                ],
                'word_count_per_section': (100, 250),
                'tone': 'persuasive',
                'structure': 'business'
            },
            'manual': {
                'sections': [
                    'Overview',
                    'Getting Started',
                    'Basic Operations',
                    'Advanced Features',
                    'Troubleshooting',
                    'FAQ',
                    'Support'
                ],
                'word_count_per_section': (100, 200),
                'tone': 'instructional',
                'structure': 'procedural'
            }
        }
    
    def _load_document_structures(self) -> Dict[str, List[str]]:
        """Load document structure patterns."""
        return {
            'academic': [
                "Abstract and Introduction",
                "Literature Review",
                "Methodology and Research Design",
                "Data Collection and Analysis",
                "Results and Discussion",
                "Conclusions and Implications",
                "Limitations and Future Research"
            ],
            'business': [
                "Executive Summary",
                "Market Analysis and Industry Trends",
                "Competitive Landscape Assessment",
                "Strategic Recommendations",
                "Financial Projections and ROI Analysis",
                "Implementation Roadmap",
                "Risk Management and Mitigation Strategies"
            ],
            'technical': [
                "Technical Overview and Architecture",
                "System Requirements and Specifications",
                "Implementation Details",
                "Performance Testing and Benchmarks",
                "Security Considerations",
                "Maintenance and Support Procedures",
                "Future Development Plans"
            ]
        }
    
    def _load_writing_styles(self) -> Dict[str, Dict]:
        """Load writing style profiles."""
        return {
            'professional': {
                'sentence_length': (15, 25),
                'paragraph_length': (3, 6),
                'vocabulary_level': 'advanced',
                'formality': 'high',
                'active_voice_ratio': 0.7
            },
            'academic': {
                'sentence_length': (20, 30),
                'paragraph_length': (4, 8),
                'vocabulary_level': 'expert',
                'formality': 'very_high',
                'active_voice_ratio': 0.6
            },
            'casual': {
                'sentence_length': (10, 20),
                'paragraph_length': (2, 4),
                'vocabulary_level': 'intermediate',
                'formality': 'low',
                'active_voice_ratio': 0.8
            }
        }
    
    def expand_prompt_10000x(self, original_prompt: str, document_type: str = "report", 
                           target_word_count: int = 5000) -> DocumentStructure:
        """Expand a simple prompt into comprehensive 10,000x content."""
        logging.info(f"Expanding prompt: '{original_prompt}' into {target_word_count}-word {document_type}")
        
        # Determine document type and structure
        if document_type not in self.content_templates:
            document_type = "report"
        
        template = self.content_templates[document_type]
        
        # Generate document structure
        structure = DocumentStructure(
            title=original_prompt.title(),
            subtitle=f"Comprehensive Analysis and Strategic Insights",
            author="JASON - AI Architect",
            date=datetime.now().strftime("%B %d, %Y"),
            sections=[],
            document_type=document_type,
            style="professional"
        )
        
        # Generate sections
        sections_template = template['sections']
        words_per_section = target_word_count // len(sections_template)
        
        for section_name in sections_template:
            section = self._generate_document_section(
                section_name, 
                original_prompt, 
                document_type, 
                words_per_section
            )
            structure.sections.append(section)
        
        # Calculate total words
        structure.total_words = sum(
            len(section.content) * 20  # Estimate words per content item
            for section in structure.sections
        )
        
        logging.info(f"Generated document structure with {len(structure.sections)} sections")
        return structure
    
    def _generate_document_section(self, section_name: str, base_topic: str, 
                                 document_type: str, word_count: int) -> DocumentSection:
        """Generate a comprehensive document section."""
        section = DocumentSection(
            title=section_name,
            content=[],
            level=1
        )
        
        # Generate content points based on section type
        content_points = self._generate_section_content_points(section_name, base_topic, document_type)
        
        # Expand each content point into detailed paragraphs
        for point in content_points:
            expanded_content = self._expand_content_point(point, base_topic, document_type)
            section.content.append(expanded_content)
        
        # Generate subsections if needed
        if section_name in ["Findings", "Analysis", "Implementation", "Methodology"]:
            subsections = self._generate_subsections(section_name, base_topic, document_type)
            section.subsections = subsections
        
        return section
    
    def _generate_section_content_points(self, section_name: str, base_topic: str, 
                                       document_type: str) -> List[str]:
        """Generate key content points for a section."""
        section_templates = {
            'Executive Summary': [
                f"Overview of {base_topic.lower()} and its strategic importance",
                f"Key findings and critical insights from comprehensive analysis",
                f"Primary recommendations and actionable next steps",
                f"Expected outcomes and impact assessment"
            ],
            'Introduction': [
                f"Background and context for {base_topic.lower()}",
                f"Problem statement and research objectives",
                f"Scope and limitations of the analysis",
                f"Methodology and approach overview"
            ],
            'Methodology': [
                f"Research design and data collection methods",
                f"Analytical frameworks and evaluation criteria",
                f"Tools and technologies utilized",
                f"Quality assurance and validation procedures"
            ],
            'Findings': [
                f"Primary discoveries and quantitative results",
                f"Qualitative insights and pattern analysis",
                f"Comparative analysis with benchmarks",
                f"Statistical significance and data interpretation"
            ],
            'Analysis': [
                f"In-depth examination of key trends and patterns",
                f"Root cause analysis and contributing factors",
                f"Implications for stakeholders and decision-making",
                f"Strategic opportunities and competitive advantages"
            ],
            'Recommendations': [
                f"Strategic priorities and implementation roadmap",
                f"Resource requirements and allocation strategies",
                f"Risk mitigation and contingency planning",
                f"Performance metrics and success indicators"
            ],
            'Conclusion': [
                f"Summary of key insights and conclusions",
                f"Implications for future research and practice",
                f"Limitations and areas for further investigation",
                f"Final recommendations and call to action"
            ]
        }
        
        return section_templates.get(section_name, [
            f"Comprehensive analysis of {section_name.lower()} in context of {base_topic.lower()}",
            f"Detailed examination of key factors and relationships",
            f"Evidence-based conclusions and recommendations",
            f"Strategic implications and future considerations"
        ])
    
    def _expand_content_point(self, point: str, base_topic: str, document_type: str) -> str:
        """Expand a content point into detailed paragraph."""
        # Context-aware content expansion
        expansion_patterns = [
            f"This comprehensive analysis reveals critical insights regarding {point.lower()}. ",
            f"Through detailed examination and systematic evaluation, we have identified several key factors that influence {point.lower()}. ",
            f"The evidence strongly suggests that {point.lower()} plays a pivotal role in shaping outcomes. ",
            f"Our findings demonstrate significant correlations between {point.lower()} and overall performance metrics. ",
            f"Strategic implementation of recommendations related to {point.lower()} can yield substantial improvements. "
        ]
        
        # Select appropriate expansion pattern
        pattern = random.choice(expansion_patterns)
        
        # Add domain-specific details
        if document_type == "report":
            return pattern + f"Data-driven analysis and stakeholder feedback confirm that {point.lower()} requires immediate attention and strategic planning."
        elif document_type == "essay":
            return pattern + f"Critical examination of {point.lower()} through multiple theoretical frameworks provides deeper understanding of its significance."
        elif document_type == "proposal":
            return pattern + f"Proposed solutions addressing {point.lower()} offer innovative approaches with measurable benefits and competitive advantages."
        else:
            return pattern + f"Further investigation and analysis of {point.lower()} will continue to inform our strategic decisions and operational improvements."
    
    def _generate_subsections(self, section_name: str, base_topic: str, 
                            document_type: str) -> List[DocumentSection]:
        """Generate subsections for complex sections."""
        subsection_templates = {
            'Findings': [
                'Quantitative Analysis',
                'Qualitative Insights',
                'Comparative Studies'
            ],
            'Analysis': [
                'Trend Analysis',
                'Impact Assessment',
                'Risk Evaluation'
            ],
            'Implementation': [
                'Phase 1: Foundation',
                'Phase 2: Development',
                'Phase 3: Deployment'
            ],
            'Methodology': [
                'Data Collection',
                'Analytical Methods',
                'Validation Procedures'
            ]
        }
        
        subsections = []
        for subsection_name in subsection_templates.get(section_name, []):
            subsection = DocumentSection(
                title=subsection_name,
                content=[f"Detailed analysis of {subsection_name.lower()} in context of {section_name.lower()}."],
                level=2
            )
            subsections.append(subsection)
        
        return subsections

class ExcelContentGenerator:
    """Advanced Excel content generation system."""
    
    def __init__(self):
        self.data_templates = self._load_data_templates()
        self.chart_templates = self._load_chart_templates()
        self.calculation_templates = self._load_calculation_templates()
    
    def _load_data_templates(self) -> Dict[str, Dict]:
        """Load templates for different types of Excel data."""
        return {
            'financial': {
                'columns': ['Date', 'Revenue', 'Expenses', 'Profit', 'Margin', 'Growth'],
                'rows': 12,  # Monthly data
                'formulas': ['=C2-B2', '=D2/B2'],
                'formatting': 'currency'
            },
            'project': {
                'columns': ['Task', 'Start Date', 'End Date', 'Duration', 'Status', 'Assigned To'],
                'rows': 20,
                'formulas': ['=DAYS(C2,B2)'],
                'formatting': 'date'
            },
            'inventory': {
                'columns': ['Product', 'SKU', 'Quantity', 'Unit Price', 'Total Value', 'Reorder Level'],
                'rows': 50,
                'formulas': ['=D2*C2'],
                'formatting': 'number'
            },
            'sales': {
                'columns': ['Region', 'Q1', 'Q2', 'Q3', 'Q4', 'Total', 'Target', 'Achievement %'],
                'rows': 10,
                'formulas': ['=SUM(B2:E2)', '=H2/G2'],
                'formatting': 'percentage'
            }
        }
    
    def _load_chart_templates(self) -> Dict[str, Dict]:
        """Load templates for different chart types."""
        return {
            'bar': {
                'title': 'Comparative Analysis',
                'x_axis': 'Categories',
                'y_axis': 'Values',
                'style': 'clustered'
            },
            'line': {
                'title': 'Trend Analysis',
                'x_axis': 'Time Period',
                'y_axis': 'Metrics',
                'style': 'smooth'
            },
            'pie': {
                'title': 'Distribution Analysis',
                'x_axis': 'Segments',
                'y_axis': 'Percentages',
                'style': 'exploded'
            }
        }
    
    def _load_calculation_templates(self) -> Dict[str, List[str]]:
        """Load templates for Excel calculations."""
        return {
            'financial': [
                '=SUM(B2:B13)',
                '=AVERAGE(B2:B13)',
                '=MAX(B2:B13)',
                '=MIN(B2:B13)',
                '=STDEV(B2:B13)'
            ],
            'statistical': [
                '=COUNT(A2:A100)',
                '=AVERAGE(B2:B100)',
                '=MEDIAN(B2:B100)',
                '=PERCENTILE(B2:B100,0.9)',
                '=CORREL(B2:B100,C2:C100)'
            ],
            'conditional': [
                '=IF(B2>1000,"High","Low")',
                '=IFERROR(B2/C2,"N/A")',
                '=IF(AND(B2>0,C2>0),B2*C2,0)',
                '=VLOOKUP(A2,Sheet2!A:B,2,FALSE)'
            ]
        }
    
    def generate_excel_workbook(self, topic: str, workbook_type: str = "financial", 
                              include_charts: bool = True) -> ExcelWorkbook:
        """Generate a comprehensive Excel workbook."""
        logging.info(f"Generating Excel workbook: '{topic}' of type '{workbook_type}'")
        
        if workbook_type not in self.data_templates:
            workbook_type = "financial"
        
        template = self.data_templates[workbook_type]
        
        # Create workbook structure
        workbook = ExcelWorkbook(
            title=f"{topic} - Analysis Dashboard",
            sheets=[],
            theme="professional",
            created_by="JASON - AI Architect"
        )
        
        # Generate main data sheet
        main_sheet = self._generate_data_sheet(topic, workbook_type, template)
        workbook.sheets.append(main_sheet)
        
        # Generate summary sheet if charts are included
        if include_charts:
            summary_sheet = self._generate_summary_sheet(topic, main_sheet, workbook_type)
            workbook.sheets.append(summary_sheet)
        
        # Generate analysis sheet
        analysis_sheet = self._generate_analysis_sheet(topic, main_sheet, workbook_type)
        workbook.sheets.append(analysis_sheet)
        
        logging.info(f"Generated Excel workbook with {len(workbook.sheets)} sheets")
        return workbook
    
    def _generate_data_sheet(self, topic: str, workbook_type: str, 
                           template: Dict) -> ExcelSheet:
        """Generate the main data sheet."""
        sheet = ExcelSheet(
            name="Data",
            data=[],
            headers=template['columns'],
            chart_type='bar' if workbook_type in ['financial', 'sales'] else 'line'
        )
        
        # Generate header row
        sheet.data.append(template['columns'])
        
        # Generate data rows
        num_rows = template['rows']
        for i in range(num_rows):
            row = self._generate_data_row(template, i)
            sheet.data.append(row)
        
        # Add calculations
        if workbook_type in self.calculation_templates:
            sheet.calculations = self._generate_calculations(workbook_type)
        
        return sheet
    
    def _generate_data_row(self, template: Dict, row_index: int) -> List[Union[str, int, float]]:
        """Generate a single data row."""
        row = []
        formatting = template.get('formatting', 'number')
        
        for col_index, column in enumerate(template['columns']):
            if column.lower() in ['date', 'start date', 'end date']:
                # Generate date
                date_value = f"2024-{(row_index % 12) + 1:02d}-{(row_index % 28) + 1:02d}"
                row.append(date_value)
            elif column.lower() in ['task', 'product', 'region']:
                # Generate categorical data
                categories = self._get_category_data(column.lower())
                row.append(categories[row_index % len(categories)])
            elif column.lower() in ['status']:
                statuses = ['Completed', 'In Progress', 'Pending', 'Not Started']
                row.append(statuses[row_index % len(statuses)])
            elif column.lower() in ['assigned to']:
                names = ['John Smith', 'Jane Doe', 'Bob Johnson', 'Alice Brown']
                row.append(names[row_index % len(names)])
            else:
                # Generate numerical data
                if formatting == 'currency':
                    value = random.randint(1000, 50000) / 100
                elif formatting == 'percentage':
                    value = random.randint(50, 150) / 100
                else:
                    value = random.randint(10, 1000)
                row.append(value)
        
        return row
    
    def _get_category_data(self, category_type: str) -> List[str]:
        """Get categorical data based on type."""
        categories = {
            'task': ['Data Analysis', 'Report Generation', 'Client Meeting', 'Code Review', 'Testing'],
            'product': ['Product A', 'Product B', 'Product C', 'Product D', 'Product E'],
            'region': ['North', 'South', 'East', 'West', 'Central']
        }
        return categories.get(category_type, ['Item A', 'Item B', 'Item C', 'Item D', 'Item E'])
    
    def _generate_calculations(self, workbook_type: str) -> List[Dict]:
        """Generate Excel calculations."""
        calculations = []
        
        if workbook_type in self.calculation_templates:
            formulas = self.calculation_templates[workbook_type]
            for i, formula in enumerate(formulas):
                calculation = {
                    'cell': f'A{100 + i}',
                    'formula': formula,
                    'description': f'Calculation {i+1}'
                }
                calculations.append(calculation)
        
        return calculations
    
    def _generate_summary_sheet(self, topic: str, data_sheet: ExcelSheet, 
                              workbook_type: str) -> ExcelSheet:
        """Generate summary sheet with charts."""
        sheet = ExcelSheet(
            name="Summary",
            data=[],
            headers=['Metric', 'Value', 'Trend'],
            chart_type='pie'
        )
        
        # Generate summary data
        summary_data = [
            ['Total Records', len(data_sheet.data) - 1, '↑'],
            ['Average Value', '=AVERAGE(Data!B2:B100)', '→'],
            ['Maximum Value', '=MAX(Data!B2:B100)', '↑'],
            ['Minimum Value', '=MIN(Data!B2:B100)', '↓'],
            ['Growth Rate', '=AVERAGE(Data!C2:C100)', '↑']
        ]
        
        sheet.data = summary_data
        
        return sheet
    
    def _generate_analysis_sheet(self, topic: str, data_sheet: ExcelSheet, 
                               workbook_type: str) -> ExcelSheet:
        """Generate analysis sheet with insights."""
        sheet = ExcelSheet(
            name="Analysis",
            data=[],
            headers=['Analysis', 'Result', 'Insights'],
            chart_type='line'
        )
        
        # Generate analysis data
        analysis_data = [
            ['Correlation Analysis', '=CORREL(Data!B2:B100,Data!C2:C100)', 'Strong positive correlation'],
            ['Trend Analysis', '=FORECAST.ETS(A101,Data!B2:B100,Data!A2:A100)', 'Upward trend expected'],
            ['Variance Analysis', '=VAR(Data!B2:B100)', 'Low variance indicates stability'],
            ['Performance Score', '=AVERAGE(Data!B2:B100)/MAX(Data!B2:B100)', 'Above average performance'],
            ['Risk Assessment', '=STDEV(Data!B2:B100)/AVERAGE(Data!B2:B100)', 'Moderate risk level']
        ]
        
        sheet.data = analysis_data
        
        return sheet

class DocumentGenerator:
    """Production-ready document and Excel generator."""
    
    def __init__(self, uspt_model=None, output_dir="/tmp"):
        self.uspt_model = uspt_model
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        self.content_expander = DocumentContentExpander()
        self.excel_generator = ExcelContentGenerator()
        
        if not DOCX_AVAILABLE:
            logging.error("python-docx not available - Word document generation disabled")
        if not OPENPYXL_AVAILABLE:
            logging.error("openpyxl not available - Excel generation disabled")
        
        logging.info(f"Document Generator initialized with output directory: {output_dir}")
    
    def create_word_document(self, prompt: str, filename: Optional[str] = None, 
                           document_type: str = "report", target_word_count: int = 5000) -> str:
        """Create a comprehensive Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word document generation")
        
        logging.info(f"Creating Word document: '{prompt}' with {target_word_count} words")
        
        # Expand prompt into comprehensive structure
        structure = self.content_expander.expand_prompt_10000x(prompt, document_type, target_word_count)
        
        # Generate filename if not provided
        if not filename:
            safe_title = re.sub(r'[^\w\s-]', '', structure.title).strip()
            safe_title = re.sub(r'[-\s]+', '_', safe_title)
            filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        output_path = self.output_dir / filename
        
        # Create Word document
        doc = Document()
        
        # Add title page
        self._add_title_page(doc, structure)
        
        # Add table of contents
        self._add_table_of_contents(doc, structure)
        
        # Add sections
        for section in structure.sections:
            self._add_document_section(doc, section)
        
        # Add footer
        self._add_footer(doc, structure)
        
        # Save document
        doc.save(str(output_path))
        
        # Save metadata
        self._save_document_metadata(structure, output_path.with_suffix('.json'))
        
        logging.info(f"Word document saved to: {output_path}")
        return str(output_path)
    
    def _add_title_page(self, doc: Document, structure: DocumentStructure):
        """Add title page to document."""
        # Add title
        title = doc.add_heading(structure.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph(structure.subtitle)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.style = 'Subtitle'
        
        # Add metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f"Author: {structure.author}\n")
        metadata.add_run(f"Date: {structure.date}\n")
        metadata.add_run(f"Document Type: {structure.document_type.title()}\n")
        metadata.add_run(f"Total Words: {structure.total_words:,}")
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document, structure: DocumentStructure):
        """Add table of contents."""
        toc_title = doc.add_heading("Table of Contents", level=1)
        
        for section in structure.sections:
            # Add section title with page number placeholder
            toc_entry = doc.add_paragraph()
            toc_entry.add_run(f"{section.title}").bold = True
            toc_entry.add_run(f"\tPage {len(doc.paragraphs) + 2}")
            
            # Add subsections if any
            for subsection in section.subsections:
                sub_toc_entry = doc.add_paragraph()
                sub_toc_entry.add_run(f"    {subsection.title}")
                sub_toc_entry.add_run(f"\tPage {len(doc.paragraphs) + 3}")
        
        doc.add_page_break()
    
    def _add_document_section(self, doc: Document, section: DocumentSection):
        """Add a section to the document."""
        # Add section heading
        heading = doc.add_heading(section.title, level=section.level)
        
        # Add section content
        for content_item in section.content:
            paragraph = doc.add_paragraph(content_item)
            paragraph.style = 'Normal'
            
            # Add some formatting for readability
            if len(content_item) > 200:
                paragraph.space_after = Pt(12)
        
        # Add subsections if any
        for subsection in section.subsections:
            self._add_document_section(doc, subsection)
        
        # Add section break for major sections
        if section.level == 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    def _add_footer(self, doc: Document, structure: DocumentStructure):
        """Add footer to document."""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated by JASON - AI Architect | {structure.title} | Page "
        footer_para.add_run("{ PAGE }").bold = True
    
    def create_excel_workbook(self, prompt: str, filename: Optional[str] = None, 
                           workbook_type: str = "financial", include_charts: bool = True) -> str:
        """Create a comprehensive Excel workbook."""
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl is required for Excel generation")
        
        logging.info(f"Creating Excel workbook: '{prompt}' of type '{workbook_type}'")
        
        # Generate workbook structure
        structure = self.excel_generator.generate_excel_workbook(prompt, workbook_type, include_charts)
        
        # Generate filename if not provided
        if not filename:
            safe_title = re.sub(r'[^\w\s-]', '', structure.title).strip()
            safe_title = re.sub(r'[-\s]+', '_', safe_title)
            filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        output_path = self.output_dir / filename
        
        # Create Excel workbook
        wb = Workbook()
        
        # Remove default sheet
        wb.remove(wb.active)
        
        # Add sheets
        for sheet_structure in structure.sheets:
            ws = wb.create_sheet(title=sheet_structure.name)
            self._add_excel_sheet(ws, sheet_structure)
        
        # Save workbook
        wb.save(str(output_path))
        
        # Save metadata
        self._save_excel_metadata(structure, output_path.with_suffix('.json'))
        
        logging.info(f"Excel workbook saved to: {output_path}")
        return str(output_path)
    
    def _add_excel_sheet(self, ws, sheet_structure: ExcelSheet):
        """Add a sheet to the Excel workbook."""
        # Add headers
        for col_idx, header in enumerate(sheet_structure.headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")
        
        # Add data
        for row_idx, row_data in enumerate(sheet_structure.data[1:], 2):  # Skip header
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                
                # Apply formatting based on data type
                if isinstance(value, str) and '$' in str(value):
                    cell.number_format = '$#,##0.00'
                elif isinstance(value, (int, float)) and col_idx > 1:
                    cell.number_format = '#,##0.00'
        
        # Add calculations
        for calc in sheet_structure.calculations:
            cell = ws[calc['cell']]
            cell.value = calc['formula']
            cell.font = Font(color="FF0000", italic=True)
        
        # Add chart if specified
        if sheet_structure.chart_type and len(sheet_structure.data) > 1:
            self._add_excel_chart(ws, sheet_structure)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    
    def _add_excel_chart(self, ws, sheet_structure: ExcelSheet):
        """Add a chart to the Excel sheet."""
        if sheet_structure.chart_type == 'bar':
            chart = BarChart()
            chart.type = "col"
            chart.style = 10
        elif sheet_structure.chart_type == 'line':
            chart = LineChart()
            chart.style = 2
        elif sheet_structure.chart_type == 'pie':
            chart = PieChart()
            chart.style = 3
        else:
            return
        
        # Set chart data
        data = Reference(ws, min_col=2, min_row=1, max_col=len(sheet_structure.headers), max_row=len(sheet_structure.data))
        cats = Reference(ws, min_col=1, min_row=2, max_row=len(sheet_structure.data))
        
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(cats)
        chart.title = f"{sheet_structure.name} Analysis"
        
        # Add chart to sheet
        ws.add_chart(chart, "G2")
    
    def _save_document_metadata(self, structure: DocumentStructure, metadata_path: Path):
        """Save document metadata to JSON file."""
        metadata = {
            'title': structure.title,
            'subtitle': structure.subtitle,
            'author': structure.author,
            'date': structure.date,
            'document_type': structure.document_type,
            'style': structure.style,
            'total_words': structure.total_words,
            'total_sections': len(structure.sections),
            'sections': [
                {
                    'title': section.title,
                    'level': section.level,
                    'content_count': len(section.content),
                    'subsections': len(section.subsections)
                }
                for section in structure.sections
            ],
            'generated_at': datetime.now().isoformat()
        }
        
        with open(metadata_path, 'w') as f:
            json.dump(metadata, f, indent=2)
        
        logging.info(f"Document metadata saved to: {metadata_path}")
    
    def _save_excel_metadata(self, structure: ExcelWorkbook, metadata_path: Path):
        """Save Excel workbook metadata to JSON file."""
        metadata = {
            'title': structure.title,
            'created_by': structure.created_by,
            'created_date': structure.created_date,
            'theme': structure.theme,
            'total_sheets': len(structure.sheets),
            'sheets': [
                {
                    'name': sheet.name,
                    'rows': len(sheet.data),
                    'columns': len(sheet.headers),
                    'chart_type': sheet.chart_type,
                    'calculations': len(sheet.calculations)
                }
                for sheet in structure.sheets
            ],
            'generated_at': datetime.now().isoformat()
        }
        
        with open(metadata_path, 'w') as f:
            json.dump(metadata, f, indent=2)
        
        logging.info(f"Excel metadata saved to: {metadata_path}")
    
    def batch_generate_documents(self, prompts: List[str], document_type: str = "report") -> List[str]:
        """Generate multiple documents from a list of prompts."""
        results = []
        
        for prompt in prompts:
            try:
                output_path = self.create_word_document(prompt, document_type=document_type)
                results.append(output_path)
                logging.info(f"Generated document for: {prompt}")
            except Exception as e:
                logging.error(f"Failed to generate document for '{prompt}': {e}")
                results.append(None)
        
        return results
    
    def batch_generate_excel_workbooks(self, prompts: List[str], workbook_type: str = "financial") -> List[str]:
        """Generate multiple Excel workbooks from a list of prompts."""
        results = []
        
        for prompt in prompts:
            try:
                output_path = self.create_excel_workbook(prompt, workbook_type=workbook_type)
                results.append(output_path)
                logging.info(f"Generated Excel workbook for: {prompt}")
            except Exception as e:
                logging.error(f"Failed to generate Excel workbook for '{prompt}': {e}")
                results.append(None)
        
        return results

# Example usage
if __name__ == '__main__':
    generator = DocumentGenerator(output_dir="/tmp")
    
    # Generate Word document
    doc_path = generator.create_word_document(
        "Q4 Business Strategy Report",
        document_type="report",
        target_word_count=5000
    )
    print(f"Word document generated: {doc_path}")
    
    # Generate Excel workbook
    excel_path = generator.create_excel_workbook(
        "Financial Analysis Dashboard",
        workbook_type="financial",
        include_charts=True
    )
    print(f"Excel workbook generated: {excel_path}")

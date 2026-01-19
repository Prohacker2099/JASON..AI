import os
import logging
import uvicorn
from fastapi import FastAPI, UploadFile, File, HTTPException, Body
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import json
import asyncio
import numpy as np
import cv2
from PIL import Image
import io
import time

# Configure Logging
import logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("JasonEngine")


from contextlib import asynccontextmanager

# Removed torch dependency for stability

# Import existing VLM module
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    from jason_service.ai_engine.vlm_gui_learner import UniversalGUILearner
except ImportError:
    logger.warning("VLM module not found.")
    UniversalGUILearner = None

    UniversalGUILearner = None

try:
    from jason_service.ai_engine.local_llm import LocalLLM
except ImportError:
    logger.warning("LocalLLM module not found.")
    LocalLLM = None

# Global Models
vlm_learner: Optional[UniversalGUILearner] = None
whisper_model = None
local_llm: Optional["LocalLLM"] = None

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup: Load models
    global vlm_learner, whisper_model
    logger.info("Initializing JASON Engine...")
    
    # Initialize VLM
    try:
        if UniversalGUILearner:
            logger.info("Loading VLM (Moondream2)...")
            # We'll initialize but skip automatic weight loading if it might hang
            vlm_learner = UniversalGUILearner()
            try:
                # Use a timeout or catch specific failure to load weights
                vlm_learner._initialize_vlm()
                logger.info("VLM Loaded Successfully.")
            except Exception as e:
                logger.error(f"VLM Weights could not be loaded: {e}. Falling back to OCR/Contour analysis.")
                vlm_learner.model = None # Ensure it handles None model correctly
        else:
            logger.warning("VLM component not available, skipping initialization.")
    except Exception as e:
        logger.error(f"Failed to load VLM: {e}")
    
    # Initialize Whisper
    try:
        import whisper
        if whisper: # Check if the module was successfully imported
            logger.info("Loading Whisper (base)...")
            whisper_model = whisper.load_model("base")
            logger.info("Whisper Loaded Successfully.")
        else:
            logger.warning("Whisper component not available, skipping initialization.")
    except ImportError:
        logger.warning("Whisper module not found, skipping initialization.")
    except Exception as e:
        logger.error(f"Failed to load Whisper: {e}")

    # Initialize Local LLM (The Reasoning Brain)
    try:
        if LocalLLM:
            logger.info("Initializing Local LLM (Reasoning Brain)...")
            local_llm = LocalLLM()
            if local_llm.available:
                logger.info(f"Local LLM Connected: {local_llm.model}")
            else:
                logger.warning("Local LLM (Ollama) not available. Brain will operate in fallback mode.")
        else:
            local_llm = None
            logger.warning("LocalLLM module missing.")
    except Exception as e:
        logger.error(f"Failed to load Local LLM: {e}")
        local_llm = None
        
    yield
    
    # Shutdown
    logger.info("Shutting down JASON Engine...")

app = FastAPI(title="JASON Engine", version="1.0.0", lifespan=lifespan)

class UIActionRequest(BaseModel):
    action_type: str
    parameters: Dict[str, Any]

class TranscriptionRequest(BaseModel):
    language: Optional[str] = "en"

class PlanningRequest(BaseModel):
    prompt: str
    context: Optional[Dict[str, Any]] = {}


@app.get("/health")
async def health_check():
    return {
        "status": "online",
        "vlm_loaded": vlm_learner is not None and vlm_learner.model is not None,
        "whisper_loaded": whisper_model is not None,
        "brain_loaded": local_llm is not None and local_llm.available
    }

@app.post("/analyze")
async def analyze_vlm(request: Dict[str, Any] = Body(...)):
    """
    Compatibility endpoint for WindowsUIAutomationAgent.
    Expects JSON body with 'image' (path) and 'prompt'.
    """
    if not vlm_learner:
        return {"ok": False, "error": "VLM not initialized"}

    image_path = request.get("image") or request.get("imagePath")
    prompt = request.get("prompt")
    
    if not image_path or not prompt:
        return {"ok": False, "error": "image and prompt are required"}

    try:
        if not os.path.exists(image_path):
            return {"ok": False, "error": f"Image not found: {image_path}"}
            
        # load image
        image = cv2.imread(image_path)
        if image is None:
            return {"ok": False, "error": "Failed to decode image"}
            
        # Use VLM to answer question
        # Note: vlm_learner has analyze_screen_semantics, but for /analyze 
        # we often just need the raw answer or coordinates.
        # We'll use the internal model if available.
        
        # Use refactored VLM learner (Ollama logic)
        analysis = vlm_learner.analyze_screen_semantics(image)
        response = analysis.layout_description

        # Extract coordinates if present
        import re
        m = re.search(r"\{[\s\S]*\}", response)
        if m:
            try:
                obj = json.loads(m.group(0))
                if "x" in obj and "y" in obj:
                    return {"ok": True, "x": float(obj["x"]), "y": float(obj["y"]), "raw": response}
            except:
                pass

        return {"ok": True, "raw": response}

    except Exception as e:
        logger.error(f"VLM Analysis failed: {e}")
        return {"ok": False, "error": str(e)}

@app.post("/analyze_screen")
async def analyze_screen(file: UploadFile = File(...)):
    """
    Analyze an uploaded screenshot using VLM.
    """
    if not vlm_learner:
        raise HTTPException(status_code=503, detail="VLM not initialized")
    
    try:
        contents = await file.read()
        nparr = np.frombuffer(contents, np.uint8)
        image = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if image is None:
            raise HTTPException(status_code=400, detail="Invalid image file")
            
        analysis = vlm_learner.analyze_screen_semantics(image)
        
        # Convert dataclass to dict for JSON response
        return {
            "ui_elements": [
                {
                    "type": e.element_type,
                    "text": e.text,
                    "bbox": e.bbox,
                    "confidence": e.confidence,
                    "semantic": e.semantic_meaning,
                    "actions": e.action_hints
                } for e in analysis.ui_elements
            ],
            "layout": analysis.layout_description,
            "confidence": analysis.confidence_score
        }
    except Exception as e:
        logger.error(f"Screen analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/transcribe")
async def transcribe_audio(file: UploadFile = File(...), language: str = "en"):
    """
    Transcribe uploaded audio file using persistent Whisper model.
    """
    if not whisper_model:
        raise HTTPException(status_code=503, detail="Whisper not initialized")
    
    temp_filename = f"temp_{file.filename}"
    try:
        contents = await file.read()
        with open(temp_filename, "wb") as f:
            f.write(contents)
            
        result = whisper_model.transcribe(temp_filename, language=language if language != "auto" else None)
        return {
            "text": result["text"].strip(),
            "confidence": 0.9 # Whisper doesn't always give overall confidence easily in simple API
        }
    except Exception as e:
        logger.error(f"Transcription failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(temp_filename):
            os.remove(temp_filename)

@app.post("/plan_task")
async def plan_task(request: PlanningRequest):
    """
    The 'Hyper-Prompt' Engine: Decomposes a natural language request into atomic UniversalActions.
    """
    if not local_llm or not local_llm.available:
        # Fallback for when LLM is offline - return a mock plan for demo/testing
        logger.warning("Local LLM offline, using simple fallback logic.")
        return {
            "plan": [
                {"type": "system", "command": f"echo 'JASON Brain is running in Offline Mode (Local LLM not connected). logic fell back to Python. prompt: {request.prompt}'"}
            ]
        }

    prompt = request.prompt
    schema = """
    {
        "plan": [
            {
                "type": "web" | "system" | "file" | "ui",
                "category": "browse" | "click" | "type" | "command" | "read" | "write",
                "url": "optional url",
                "command": "optional system command",
                "selector": "optional css selector",
                "value": "optional text value or file path",
                "description": "what this step does"
            }
        ]
    }
    """
    
    full_prompt = f"""
    You are JASON's Reasoning Brain. Your goal is to convert the user's HIGH-LEVEL prompt into a list of EXECUTABLE STEPS (actions).
    
    User Prompt: "{prompt}"
    
    Available Tool Capabilities:
    1. Web: browse, click, type, extract (scrape)
    2. System: Run terminal commands (e.g. 'calc', 'notepad', 'shutdown')
    3. File: Read/Write files
    4. UI: Click/Type on screen coordinates (handled via 'ui' type)
    
    Return ONLY JSON matching the schema.
    """
    
    try:
        plan = local_llm.generate_json(full_prompt, schema)
        if not plan:
            raise HTTPException(status_code=500, detail="Failed to generate plan from LLM")
        return plan
    except Exception as e:
        logger.error(f"Planning failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/execute_ui_action")
async def execute_ui_action(request: UIActionRequest):
    """
    Execute low-level UI actions (mouse/keyboard) via pynput/pyautogui.
    This runs on the host (or hidden desktop where this script is running).
    """
    try:
        import pyautogui
        
        action = request.action_type
        params = request.parameters
        
        if action == "click":
            x = params.get("x")
            y = params.get("y")
            clicks = params.get("clicks", 1)
            button = params.get("button", "left")
            if x is not None and y is not None:
                pyautogui.click(x=x, y=y, clicks=clicks, button=button)
                
        elif action == "type":
            text = params.get("text", "")
            interval = params.get("interval", 0.0)
            pyautogui.write(text, interval=interval)
            
        elif action == "press":
            keys = params.get("keys", [])
            if isinstance(keys, str):
                keys = [keys]
            for key in keys:
                pyautogui.press(key)
                
        elif action == "scroll":
            clicks = params.get("clicks", 0)
            pyautogui.scroll(clicks)
            
        else:
            raise HTTPException(status_code=400, detail=f"Unknown action: {action}")
            
        return {"status": "success", "action": action}
        
    except Exception as e:
        logger.error(f"UI Action failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class ReportRequest(BaseModel):
    title: str
    sections: List[Dict[str, str]]
    output_path: Optional[str] = None

@app.post("/generate_report")
async def generate_report(request: ReportRequest):
    """
    Generate a Word document (.docx) from structured data.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        # Title
        doc.add_heading(request.title, 0)
        
        # Sections
        for section in request.sections:
            if "heading" in section:
                doc.add_heading(section["heading"], level=1)
            if "content" in section:
                doc.add_paragraph(section["content"])
            if "list" in section:
                # Assuming simple string list for now
                if isinstance(section["list"], list):
                    for item in section["list"]:
                        doc.add_paragraph(str(item), style='List Bullet')

        # Save
        filename = request.output_path or f"report_{int(time.time())}.docx"
        # Ensure it saves in a user-accessible place if not specified
        if not os.path.isabs(filename):
            # Save to downloads or desktop by default in a real app, 
            # here we'll save to CWD (project root usually)
            filename = os.path.abspath(filename)
            
        doc.save(filename)
        
        return {"status": "success", "filename": filename}
        
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx not installed")
    except Exception as e:
        logger.error(f"Report generation failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    # In production, this might be launched via ghost workspace manager on a specific port/desktop
    uvicorn.run(app, host="0.0.0.0", port=8000)
